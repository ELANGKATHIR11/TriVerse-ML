"""
DOCX Report Generator for CodeAlpha Enterprise AI Platform.

Generates a styled Microsoft Word document report using python-docx,
mirroring the structure of the PDF report with all sections, tables,
and embedded SHAP images.
"""

from __future__ import annotations

import base64
import logging
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Brand colours (DOCX uses RGB tuples)
# ---------------------------------------------------------------------------

_PRIMARY_RGB = RGBColor(0x4F, 0x46, 0xE5)     # Indigo 600
_SECONDARY_RGB = RGBColor(0x7C, 0x3A, 0xED)   # Violet 600
_HEADER_RGB = RGBColor(0x31, 0x2E, 0x81)      # Indigo 900
_ACCENT_RGB = RGBColor(0x06, 0xB6, 0xD4)      # Cyan 500
_TEXT_RGB = RGBColor(0x1E, 0x1B, 0x4B)        # Dark Indigo
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
_ALT_ROW_HEX = "F5F3FF"
_HEADER_HEX = "312E81"
_PRIMARY_HEX = "4F46E5"
_SECONDARY_HEX = "7C3AED"


# ---------------------------------------------------------------------------
# DOCX XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set a table cell background colour using OOXML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    color: RGBColor | None = None,
    font_size: int = 9,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    """Set cell text with formatting."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------

class DOCXReportGenerator:
    """
    Generates a styled Microsoft Word (.docx) experiment report.

    Usage::

        gen = DOCXReportGenerator()
        path = gen.generate(output_path=Path("report.docx"), ...)
    """

    def generate(
        self,
        output_path: Path,
        title: str,
        task_type: str,
        dataset_summary: dict[str, Any],
        preprocessing_steps: list[str],
        metrics_table: list[dict[str, Any]],
        leaderboard: list[dict[str, Any]],
        shap_images: dict[str, str],
        ai_insights: str,
        recommendations: list[str],
    ) -> Path:
        """
        Build and save the Word document.

        Args:
            output_path:         Destination ``.docx`` file path.
            title:               Report title.
            task_type:           ML task type.
            dataset_summary:     Dataset metadata dict.
            preprocessing_steps: Ordered preprocessing descriptions.
            metrics_table:       Per-model evaluation metric rows.
            leaderboard:         Ranked leaderboard entries.
            shap_images:         Base64-encoded SHAP PNG per model.
            ai_insights:         AI-generated narrative text.
            recommendations:     Action recommendation strings.

        Returns:
            The ``output_path`` after saving.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._configure_document(doc)

        # ----------------------------------------------------------------
        # Cover page
        # ----------------------------------------------------------------
        self._add_cover(doc, title, task_type)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # Table of contents placeholder
        # ----------------------------------------------------------------
        self._add_heading(doc, "Table of Contents", level=1)
        toc_sections = [
            "1. Executive Summary",
            "2. Dataset Overview",
            "3. Preprocessing Pipeline",
            "4. Performance Metrics",
            "5. Leaderboard",
            "6. Explainability (SHAP)",
            "7. AI Insights",
            "8. Recommendations",
            "9. Future Work",
        ]
        for entry in toc_sections:
            p = doc.add_paragraph(entry)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.size = Pt(10)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 1. Executive Summary
        # ----------------------------------------------------------------
        self._add_heading(doc, "1. Executive Summary", level=1)
        best = leaderboard[0] if leaderboard else {}
        best_name = best.get("model_name", best.get("model", "N/A"))
        best_score = best.get("accuracy", best.get("weighted_score", "N/A"))

        exec_text = (
            f"This report evaluates {len(metrics_table)} machine learning models "
            f"for a {task_type} task. "
            f"The top-performing model is {best_name} with a score of {best_score}. "
            f"The dataset contains {dataset_summary.get('total_rows', 'N/A')} samples "
            f"and {dataset_summary.get('total_cols', 'N/A')} features. "
            "All models were trained and evaluated by the CodeAlpha AutoML pipeline."
        )
        doc.add_paragraph(exec_text)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 2. Dataset Overview
        # ----------------------------------------------------------------
        self._add_heading(doc, "2. Dataset Overview", level=1)
        ds_rows = [(k.replace("_", " ").title(), str(v)) for k, v in dataset_summary.items()]
        if ds_rows:
            table = doc.add_table(rows=1 + len(ds_rows), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            # Header
            _set_cell_bg(table.rows[0].cells[0], _PRIMARY_HEX)
            _set_cell_bg(table.rows[0].cells[1], _PRIMARY_HEX)
            _set_cell_text(table.rows[0].cells[0], "Property", bold=True, color=_WHITE_RGB, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(table.rows[0].cells[1], "Value", bold=True, color=_WHITE_RGB, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            for i, (key, val) in enumerate(ds_rows, start=1):
                if i % 2 == 0:
                    _set_cell_bg(table.rows[i].cells[0], _ALT_ROW_HEX)
                    _set_cell_bg(table.rows[i].cells[1], _ALT_ROW_HEX)
                _set_cell_text(table.rows[i].cells[0], key, align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell_text(table.rows[i].cells[1], val, align=WD_ALIGN_PARAGRAPH.LEFT)
            table.columns[0].width = Inches(2.5)
            table.columns[1].width = Inches(3.5)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 3. Preprocessing Pipeline
        # ----------------------------------------------------------------
        self._add_heading(doc, "3. Preprocessing Pipeline", level=1)
        doc.add_paragraph(
            "The following preprocessing steps were applied to the raw dataset "
            "before model training:"
        )
        for i, step in enumerate(preprocessing_steps, start=1):
            p = doc.add_paragraph(f"Step {i}: {step}", style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 4. Performance Metrics
        # ----------------------------------------------------------------
        self._add_heading(doc, "4. Performance Metrics", level=1)
        if metrics_table:
            keys = list(metrics_table[0].keys())
            table = doc.add_table(rows=1 + len(metrics_table), cols=len(keys))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            # Header row
            for j, col_name in enumerate(keys):
                hdr = table.rows[0].cells[j]
                _set_cell_bg(hdr, _HEADER_HEX)
                _set_cell_text(
                    hdr,
                    col_name.replace("_", " ").title(),
                    bold=True, color=_WHITE_RGB, font_size=9,
                )
            # Data rows
            for i, row_data in enumerate(metrics_table, start=1):
                bg = _ALT_ROW_HEX if i % 2 == 0 else "FFFFFF"
                for j, key in enumerate(keys):
                    cell = table.rows[i].cells[j]
                    if bg != "FFFFFF":
                        _set_cell_bg(cell, bg)
                    val = row_data.get(key, "—")
                    if isinstance(val, float):
                        val = f"{val:.4f}"
                    _set_cell_text(cell, str(val), font_size=9)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 5. Leaderboard
        # ----------------------------------------------------------------
        self._add_heading(doc, "5. Leaderboard", level=1)
        doc.add_paragraph(
            "Models ranked by weighted composite score: "
            "Accuracy×0.4 + Precision×0.2 + Recall×0.2 + "
            "(1/inference_ms)×0.1 + (1/training_sec)×0.1"
        )
        if leaderboard:
            lb_keys = list(leaderboard[0].keys())
            table = doc.add_table(rows=1 + len(leaderboard), cols=len(lb_keys))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for j, col_name in enumerate(lb_keys):
                hdr = table.rows[0].cells[j]
                _set_cell_bg(hdr, _PRIMARY_HEX)
                _set_cell_text(
                    hdr, col_name.replace("_", " ").title(),
                    bold=True, color=_WHITE_RGB, font_size=9,
                )
            for i, entry in enumerate(leaderboard, start=1):
                is_top = i == 1
                bg = "FEF9C3" if is_top else (_ALT_ROW_HEX if i % 2 == 0 else "FFFFFF")
                for j, key in enumerate(lb_keys):
                    cell = table.rows[i].cells[j]
                    _set_cell_bg(cell, bg)
                    val = entry.get(key, "—")
                    if isinstance(val, float):
                        val = f"{val:.4f}"
                    _set_cell_text(cell, str(val), bold=is_top, font_size=9)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 6. SHAP Explainability
        # ----------------------------------------------------------------
        self._add_heading(doc, "6. Explainability (SHAP)", level=1)
        doc.add_paragraph(
            "SHAP (SHapley Additive exPlanations) values show feature contribution "
            "to individual predictions. Positive values increase, negative decrease predictions."
        )
        for model_name, b64_png in shap_images.items():
            self._add_heading(doc, f"Model: {model_name}", level=2)
            try:
                raw = base64.b64decode(b64_png)
                img_stream = BytesIO(raw)
                doc.add_picture(img_stream, width=Inches(5.5))
                caption = doc.add_paragraph(f"Figure: SHAP summary plot for {model_name}.")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.font.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            except Exception as exc:
                logger.warning("Could not embed SHAP image for %s: %s", model_name, exc)
                doc.add_paragraph("[SHAP image unavailable]")
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 7. AI Insights
        # ----------------------------------------------------------------
        self._add_heading(doc, "7. AI Insights", level=1)
        doc.add_paragraph(
            "The following insights were generated by the CodeAlpha AI assistant "
            "using Retrieval-Augmented Generation (RAG):"
        )
        for line in ai_insights.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("•") or line.startswith("-"):
                p = doc.add_paragraph(line.lstrip("•- "), style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.3)
            else:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(10)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 8. Recommendations
        # ----------------------------------------------------------------
        self._add_heading(doc, "8. Recommendations", level=1)
        for i, rec in enumerate(recommendations, start=1):
            p = doc.add_paragraph(f"{i}. {rec}", style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
        doc.add_page_break()

        # ----------------------------------------------------------------
        # 9. Future Work
        # ----------------------------------------------------------------
        self._add_heading(doc, "9. Future Work", level=1)
        future_items = [
            "Expand dataset with additional labelled samples.",
            "Explore ensemble stacking of the top-3 leaderboard models.",
            "Apply neural architecture search (NAS) for deep models.",
            "Implement online learning for continuous retraining.",
            "Add fairness metrics for regulated domains.",
            "Profile memory and compute for edge deployment.",
        ]
        for item in future_items:
            p = doc.add_paragraph(item, style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)

        # ----------------------------------------------------------------
        # Save
        # ----------------------------------------------------------------
        doc.save(str(output_path))
        logger.info("DOCX report written to %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _configure_document(self, doc: Document) -> None:
        """Set default document font and margins."""
        from docx.oxml.ns import qn as _qn
        section = doc.sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        # Default paragraph font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = _TEXT_RGB

    def _add_cover(self, doc: Document, title: str, task_type: str) -> None:
        """Add a styled cover page."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CodeAlpha")
        run.bold = True
        run.font.size = Pt(32)
        run.font.color.rgb = _PRIMARY_RGB

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Enterprise AI Platform")
        run2.font.size = Pt(16)
        run2.font.color.rgb = _SECONDARY_RGB

        doc.add_paragraph()

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(title)
        run3.bold = True
        run3.font.size = Pt(20)
        run3.font.color.rgb = _TEXT_RGB

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run(f"Task Type: {task_type.title()}")
        run4.font.size = Pt(12)
        run4.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)

        for _ in range(4):
            doc.add_paragraph()

        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = p5.add_run("Generated by CodeAlpha AutoML Pipeline")
        run5.italic = True
        run5.font.size = Pt(9)
        run5.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        """Add a coloured heading paragraph."""
        heading_sizes = {1: Pt(14), 2: Pt(12), 3: Pt(11)}
        heading_colors = {
            1: _PRIMARY_RGB,
            2: _SECONDARY_RGB,
            3: RGBColor(0x06, 0xB6, 0xD4),
        }
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = heading_sizes.get(level, Pt(12))
        run.font.color.rgb = heading_colors.get(level, _PRIMARY_RGB)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
